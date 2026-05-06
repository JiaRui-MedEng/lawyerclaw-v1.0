"""
文档解析器
支持 PDF 和 Word 文档自动识别和解析
"""
import os
import logging
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime

logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器 - 支持 PDF 和 Word"""
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.doc', '.txt'}
    
    def parse_file(self, file_path: str) -> Dict:
        """
        解析文件（自动识别格式）
        
        Args:
            file_path: 文件路径
        
        Returns:
            {
                'file_path': str,
                'file_name': str,
                'file_type': str,  # 'pdf' | 'word'
                'content': str,
                'chunks': List[str],
                'metadata': Dict
            }
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在：{file_path}")
        
        # 自动识别文件类型
        file_type = self._detect_file_type(file_path)
        logger.info(f"📄 解析文件：{file_path.name} (类型：{file_type})")
        
        # 根据类型解析
        if file_type == 'pdf':
            content = self._parse_pdf(file_path)
        elif file_type == 'word':
            content = self._parse_word(file_path)
        elif file_type == 'txt':
            content = self._parse_txt(file_path)
        else:
            raise ValueError(f"不支持的文件类型：{file_path.suffix}")
        
        # 分块
        chunks = self._chunk_text(content)
        
        # 提取元数据
        metadata = self._extract_metadata(file_path, file_type)
        
        return {
            'file_path': str(file_path),
            'file_name': file_path.name,
            'file_type': file_type,
            'content': content,
            'chunks': chunks,
            'metadata': metadata,
            'parsed_at': datetime.now().isoformat()
        }
    
    def _detect_file_type(self, file_path: Path) -> str:
        """检测文件类型"""
        suffix = file_path.suffix.lower()

        if suffix == '.pdf':
            return 'pdf'
        elif suffix in ['.docx', '.doc']:
            return 'word'
        elif suffix == '.txt':
            return 'txt'
        else:
            raise ValueError(f"不支持的文件格式：{suffix}")
    
    def _parse_pdf(self, file_path: Path) -> str:
        """解析 PDF 文件"""
        try:
            import pdfplumber
            
            content_parts = []
            
            with pdfplumber.open(file_path) as pdf:
                logger.info(f"  PDF 页数：{len(pdf.pages)}")
                
                for i, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"[第 {i} 页]\n{text}")
            
            content = "\n\n".join(content_parts)
            logger.info(f"✅ PDF 解析成功，总字符数：{len(content)}")
            
            return content
            
        except ImportError:
            logger.error("❌ 未安装 pdfplumber，请运行：pip install pdfplumber")
            raise
        except Exception as e:
            logger.error(f"❌ PDF 解析失败：{e}")
            raise
    
    def _parse_word(self, file_path: Path) -> str:
        """解析 Word 文件"""
        try:
            from docx import Document

            doc = Document(file_path)

            content_parts = []

            # 提取段落
            for i, para in enumerate(doc.paragraphs, 1):
                if para.text.strip():
                    content_parts.append(para.text)

            # 提取表格
            for i, table in enumerate(doc.tables, 1):
                content_parts.append(f"\n[表格 {i}]\n")
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    content_parts.append(" | ".join(row_text))

            content = "\n\n".join(content_parts)
            logger.info(f"✅ Word 解析成功，段落数：{len(doc.paragraphs)}, 表格数：{len(doc.tables)}")
            logger.info(f"   总字符数：{len(content)}")

            return content

        except ImportError:
            logger.error("❌ 未安装 python-docx，请运行：pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"❌ Word 解析失败：{e}")
            raise

    def _parse_txt(self, file_path: Path) -> str:
        """解析 TXT 文件"""
        try:
            content = file_path.read_text(encoding='utf-8')
            logger.info(f"✅ TXT 解析成功，总字符数：{len(content)}")
            return content
        except UnicodeDecodeError:
            # 尝试 GBK 编码
            content = file_path.read_text(encoding='gbk')
            logger.info(f"✅ TXT 解析成功 (GBK 编码)，总字符数：{len(content)}")
            return content
        except Exception as e:
            logger.error(f"❌ TXT 解析失败：{e}")
            raise
    
    def _chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[Dict]:
        """
        文本分块 - 法律文档智能分块

        优先按法条结构分块（第X条），回退到固定大小分块。
        自动过滤纯目录内容。
        返回结构化 chunk 列表。

        Args:
            text: 输入文本
            chunk_size: 每块大小（字符数）
            overlap: 重叠部分大小

        Returns:
            结构化分块列表: [{'content': str, 'article_number': str, 'chapter': str, 'law_name': str}, ...]
        """
        if len(text) <= chunk_size:
            return [{'content': text, 'article_number': '', 'chapter': '', 'law_name': self._extract_law_name(text)}]

        # 尝试法律文档专用分块（返回结构化结果）
        legal_chunks = self._chunk_legal_text(text, chunk_size)
        if legal_chunks:
            logger.info(f"📝 法律文档智能分块完成，共 {len(legal_chunks)} 块")
            return legal_chunks

        # 回退到通用分块
        raw_chunks = self._chunk_generic(text, chunk_size, overlap)
        law_name = self._extract_law_name(text)
        chunks = [{'content': c, 'article_number': '', 'chapter': '', 'law_name': law_name} for c in raw_chunks]
        logger.info(f"📝 文本分块完成，共 {len(chunks)} 块，每块约 {chunk_size} 字符")
        return chunks

    def _chunk_legal_text(self, text: str, chunk_size: int = 500) -> Optional[List[Dict]]:
        """
        法律文档专用分块：按"第X条"切分，保证每条法条完整
        返回结构化 chunk，包含法条编号、章节、法律名称

        策略：
        1. 去除纯目录部分
        2. 按"第X条"分割，每条法条完整保留（绝不拆分）
        3. 短法条就近合并，直到接近 chunk_size
        4. 超长法条单独成块，不做截断
        5. 跨章节时强制切分
        """
        import re

        # 检测是否为法律文档（包含"第X条"模式）
        article_pattern = re.compile(r'(第[一二三四五六七八九十百千零０-９\d]+条)')
        matches = list(article_pattern.finditer(text))

        if len(matches) < 5:
            return None

        # 提取法律名称（从文档前 200 字符中提取）
        law_name = self._extract_law_name(text)

        # 去除目录部分
        content_start = 0
        toc_pattern = re.compile(r'目\s*录')
        toc_match = toc_pattern.search(text[:2000])

        if toc_match:
            for m in matches:
                after_text = text[m.end():m.end() + 100]
                if len(after_text.split('\n')[0].strip()) > 15:
                    content_start = m.start()
                    break

        body_text = text[content_start:]

        # 按"第X条"分割，提取每条完整法条及其编号
        article_starts = list(article_pattern.finditer(body_text))
        articles = []  # [(article_text, article_number)]
        for i, match in enumerate(article_starts):
            start = match.start()
            end = article_starts[i + 1].start() if i + 1 < len(article_starts) else len(body_text)
            article_text = body_text[start:end].strip()
            if article_text:
                articles.append((article_text, match.group(1)))

        if not articles:
            return None

        # 检测章节边界并跟踪当前章节
        chapter_pattern = re.compile(r'(第[一二三四五六七八九十百千]+(?:编|分编|章|节))\s*(.*)')
        current_chapter = ""

        # 在正文中预扫描章节标记位置
        chapter_markers = {}  # article_index -> chapter_name
        for i, (article_text, _) in enumerate(articles):
            # 检查法条前面的文本是否包含章节标记
            cm = chapter_pattern.match(article_text)
            if cm:
                current_chapter = f"{cm.group(1)} {cm.group(2)}".strip()
                chapter_markers[i] = current_chapter

        # 也扫描法条之间的章节标记
        current_chapter = ""
        full_text_chapters = list(chapter_pattern.finditer(body_text))
        chapter_positions = [(m.start(), f"{m.group(1)} {m.group(2)}".strip()) for m in full_text_chapters]

        # 为每条法条分配章节
        article_chapters = []
        ch_idx = 0
        for i, (article_text, _) in enumerate(articles):
            # 找到该法条之前最近的章节标记
            article_pos = body_text.find(article_text)
            while ch_idx < len(chapter_positions) - 1 and chapter_positions[ch_idx + 1][0] <= article_pos:
                ch_idx += 1
            if ch_idx < len(chapter_positions) and chapter_positions[ch_idx][0] <= article_pos:
                article_chapters.append(chapter_positions[ch_idx][1])
            else:
                article_chapters.append("")

        # 合并短法条，保证每条完整，跨章节时强制切分
        chunks = []
        buffer_texts = []
        buffer_articles = []  # 记录 buffer 中的法条编号
        buffer_chapter = ""
        buffer_len = 0

        for i, (article_text, article_num) in enumerate(articles):
            chapter = article_chapters[i] if i < len(article_chapters) else ""

            # 跨章节强制切分
            if buffer_texts and chapter and chapter != buffer_chapter:
                chunks.append({
                    'content': "\n".join(buffer_texts).strip(),
                    'article_number': self._format_article_range(buffer_articles),
                    'chapter': buffer_chapter,
                    'law_name': law_name,
                })
                buffer_texts = []
                buffer_articles = []
                buffer_len = 0

            buffer_chapter = chapter or buffer_chapter

            new_len = buffer_len + len(article_text) + 1
            if buffer_texts and new_len > chunk_size:
                # buffer 满了，存入并开始新 buffer
                chunks.append({
                    'content': "\n".join(buffer_texts).strip(),
                    'article_number': self._format_article_range(buffer_articles),
                    'chapter': buffer_chapter,
                    'law_name': law_name,
                })
                buffer_texts = [article_text]
                buffer_articles = [article_num]
                buffer_len = len(article_text)
            else:
                buffer_texts.append(article_text)
                buffer_articles.append(article_num)
                buffer_len = new_len

        # 处理剩余 buffer
        if buffer_texts:
            chunks.append({
                'content': "\n".join(buffer_texts).strip(),
                'article_number': self._format_article_range(buffer_articles),
                'chapter': buffer_chapter,
                'law_name': law_name,
            })

        return chunks if chunks else None

    @staticmethod
    def _extract_law_name(text: str) -> str:
        """从文档前部提取法律名称"""
        import re
        # 常见法律名称模式：《XXX》或 "中华人民共和国XXX法/条例"
        bracket_match = re.search(r'《([^》]{2,30})》', text[:500])
        if bracket_match:
            return bracket_match.group(1)

        law_match = re.search(r'(中华人民共和国[^\n]{2,20}(?:法|条例|规定|办法))', text[:500])
        if law_match:
            return law_match.group(1)

        # 取第一行非空文本作为法律名称
        for line in text[:300].split('\n'):
            line = line.strip()
            if len(line) > 4 and len(line) < 50:
                return line

        return ""

    @staticmethod
    def _format_article_range(articles: List[str]) -> str:
        """格式化法条编号范围，如 ['第一条', '第二条', '第三条'] -> '第一条-第三条'"""
        if not articles:
            return ""
        if len(articles) == 1:
            return articles[0]
        return f"{articles[0]}-{articles[-1]}"

    def _chunk_generic(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """通用分块（固定大小 + 句子边界）"""
        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # 尝试在句子边界处切分
            if end < len(text):
                for sep in ['。\n', '\n\n', '。', '\n']:
                    last_sep = text[start:end].rfind(sep)
                    if last_sep > 0:
                        end = start + last_sep + len(sep)
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - overlap

        return chunks
    
    def _extract_metadata(self, file_path: Path, file_type: str) -> Dict:
        """提取文件元数据"""
        stat = file_path.stat()
        
        return {
            'file_size': stat.st_size,
            'file_type': file_type,
            'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'absolute_path': str(file_path.absolute())
        }
    
    def parse_directory(self, directory_path: str, recursive: bool = True) -> List[Dict]:
        """
        解析目录下所有文档
        
        Args:
            directory_path: 目录路径
            recursive: 是否递归子目录
        
        Returns:
            解析结果列表
        """
        directory = Path(directory_path)
        
        if not directory.exists():
            raise FileNotFoundError(f"目录不存在：{directory}")
        
        logger.info(f"📂 解析目录：{directory}")
        
        documents = []
        
        # 遍历文件
        pattern = '**/*' if recursive else '*'
        for file_path in directory.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    doc_data = self.parse_file(file_path)
                    documents.append(doc_data)
                    logger.info(f"✅ 解析成功：{file_path.name}")
                except Exception as e:
                    logger.error(f"❌ 解析失败 {file_path.name}: {e}")
        
        logger.info(f"📊 目录解析完成，成功 {len(documents)} 个文件")
        
        return documents


# 便捷函数
def parse_document(file_path: str) -> Dict:
    """便捷函数：解析单个文档"""
    parser = DocumentParser()
    return parser.parse_file(file_path)

def parse_documents(directory: str, recursive: bool = True) -> List[Dict]:
    """便捷函数：解析目录下所有文档"""
    parser = DocumentParser()
    return parser.parse_directory(directory, recursive)
