"""
风格分析器 - 从文档中提取个人写作风格

纯本地分析，零 LLM 调用，保护隐私且零成本。

功能：
- PDF/Word/纯文本解析
- 词汇指纹提取（滑动窗口分词 + 词频统计）
- 句式节奏分析（平均句长、短句占比）
- AI 套话检测（规则匹配）
- 反 AI 规则生成
"""
import re
import logging
import statistics
from collections import Counter
from pathlib import Path
from typing import Dict, List, Optional, Any

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════
# 文档解析
# ═══════════════════════════════════════════════════════════

def parse_pdf(file_path: str) -> str:
    """
    解析 PDF 文件，提取文本内容
    
    Args:
        file_path: PDF 文件路径
        
    Returns:
        提取的文本内容
        
    Raises:
        ImportError: PyMuPDF 未安装
        FileNotFoundError: 文件不存在
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError(
            "PyMuPDF 未安装，无法解析 PDF。运行: pip install pymupdf"
        )
    
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF 文件不存在: {file_path}")
    
    doc = fitz.open(str(path))
    text_parts = []
    for page in doc:
        page_text = page.get_text()
        if page_text.strip():
            text_parts.append(page_text)
    doc.close()
    
    text = "\n".join(text_parts)
    logger.info(f"PDF 解析完成: {path.name}, {len(text)} 字符, {len(doc)} 页")
    return text


def parse_docx(file_path: str) -> str:
    """
    解析 Word 文件，提取文本内容
    
    Args:
        file_path: Word 文件路径
        
    Returns:
        提取的文本内容
        
    Raises:
        ImportError: python-docx 未安装
        FileNotFoundError: 文件不存在
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx 未安装，无法解析 Word。运行: pip install python-docx"
        )
    
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Word 文件不存在: {file_path}")
    
    doc = Document(str(path))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    text = "\n\n".join(text_parts)
    logger.info(f"Word 解析完成: {path.name}, {len(text)} 字符, {len(doc.paragraphs)} 段落")
    return text


def parse_text(file_path: str) -> str:
    """
    读取纯文本文件
    
    Args:
        file_path: 文本文件路径
        
    Returns:
        文本内容
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文本文件不存在: {file_path}")
    
    text = path.read_text(encoding="utf-8")
    logger.info(f"文本读取完成: {path.name}, {len(text)} 字符")
    return text


def parse_document(file_path: str) -> str:
    """
    自动检测文件类型并解析
    
    Args:
        file_path: 文件路径
        
    Returns:
        提取的文本内容
    """
    path = Path(file_path)
    ext = path.suffix.lower()
    
    parsers = {
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".doc": parse_docx,
        ".txt": parse_text,
        ".md": parse_text,
    }
    
    if ext not in parsers:
        raise ValueError(f"不支持的文件格式: {ext}。支持: {', '.join(parsers.keys())}")
    
    return parsers[ext](file_path)


# ═══════════════════════════════════════════════════════════
# 风格分析
# ═══════════════════════════════════════════════════════════

# 默认停用词表
DEFAULT_STOP_WORDS = {
    "的", "了", "在", "是", "我", "有", "和", "就", "不", "人", "都", "一",
    "一个", "上", "也", "很", "到", "说", "要", "去", "你", "会", "着",
    "没有", "看", "好", "自己", "这", "那", "里", "吗", "什么", "怎么",
    "这个", "那个", "可以", "能够", "需要", "应该", "我们", "他们",
}

# AI 套话规则表
AI_PATTERNS = [
    "首先", "其次", "最后", "综上所述", "总而言之",
    "值得注意的是", "需要指出的是", "不可否认",
    "不仅.*而且", "在.*方面", "对于.*来说",
    "赋能", "闭环", "抓手", "颗粒度", "对齐", "拉齐",
    "底层逻辑", "顶层设计", "方法论", "最佳实践",
    "一方面.*另一方面", "综上所述", "总而言之",
]

# 口语词 / 书面语词
COLLOQUIAL_WORDS = {
    "其实", "说白了", "挺", "蛮", "挺好", "还行", "差不多", "大概",
    "嗯", "啊", "哦", "哈哈", "嘿嘿", "然后", "不过", "话说回来",
    "我觉得", "说实话", "老实说", "反正", "毕竟",
}

FORMAL_WORDS = {
    "综上所述", "由此可见", "诚然", "毋庸置疑", "显而易见",
    "鉴于", "据此", "特此", "兹", " aforementioned",
}


def _tokenize(text: str) -> List[str]:
    """
    滑动窗口分词（2-3 字）
    
    无需 jieba，轻量快速。
    """
    words = []
    # 按标点分割为句子
    sentences = re.split(r'[。，！？；\n\r\t]', text)
    for sent in sentences:
        sent = sent.strip()
        if len(sent) < 2:
            continue
        # 滑动窗口提取 2-3 字词
        for i in range(len(sent) - 1):
            if i + 2 <= len(sent):
                words.append(sent[i:i+2])
            if i + 3 <= len(sent):
                words.append(sent[i:i+3])
    return words


def _split_sentences(text: str) -> List[str]:
    """按标点分句"""
    sentences = re.split(r'[。！？；]', text)
    return [s.strip() for s in sentences if s.strip()]


def _split_paragraphs(text: str) -> List[str]:
    """按换行分段"""
    paragraphs = re.split(r'\n\s*\n', text)
    return [p.strip() for p in paragraphs if p.strip()]


def analyze_vocabulary(text: str, stop_words: Optional[set] = None, top_n: int = 10) -> Dict[str, Any]:
    """
    词汇指纹分析
    
    Args:
        text: 输入文本
        stop_words: 停用词集合
        top_n: 返回前 N 个高频词
        
    Returns:
        词汇特征字典
    """
    if stop_words is None:
        stop_words = DEFAULT_STOP_WORDS
    
    words = _tokenize(text)
    filtered = [w for w in words if w not in stop_words and len(w) >= 2]
    
    word_freq = Counter(filtered)
    top_words = [w for w, c in word_freq.most_common(top_n)]
    
    # 计算正式度（口语词/总词数）
    colloquial_count = sum(1 for w in top_words if w in COLLOQUIAL_WORDS)
    formal_count = sum(1 for w in top_words if w in FORMAL_WORDS)
    total = colloquial_count + formal_count
    formality_score = round(formal_count / total, 2) if total > 0 else 0.5
    
    # 检测个人标志性词汇（排除常见词）
    common_words = {"这个", "那个", "什么", "怎么", "可以", "需要", "我们", "他们"}
    personal_markers = [w for w in top_words if w not in common_words and len(w) >= 2]
    
    return {
        "favorite_words": personal_markers[:5],
        "top_words": top_words,
        "word_freq": dict(word_freq.most_common(20)),
        "formality_score": formality_score,
        "total_unique_words": len(set(filtered)),
        "total_words": len(filtered),
    }


def analyze_rhythm(text: str) -> Dict[str, Any]:
    """
    句式节奏分析
    
    Args:
        text: 输入文本
        
    Returns:
        节奏特征字典
    """
    sentences = _split_sentences(text)
    if not sentences:
        return {
            "sentence_length_avg": 0,
            "sentence_length_std": 0,
            "short_sentence_ratio": 0,
            "medium_sentence_ratio": 0,
            "long_sentence_ratio": 0,
            "question_ratio": 0,
            "exclamation_ratio": 0,
        }
    
    lengths = [len(s) for s in sentences]
    avg_length = round(statistics.mean(lengths), 1)
    std_length = round(statistics.stdev(lengths), 1) if len(lengths) > 1 else 0
    
    short_count = sum(1 for l in lengths if l < 10)
    medium_count = sum(1 for l in lengths if 10 <= l <= 25)
    long_count = sum(1 for l in lengths if l > 25)
    total = len(lengths)
    
    question_count = sum(1 for s in sentences if "？" in s or "?" in s)
    exclamation_count = sum(1 for s in sentences if "！" in s or "!" in s)
    
    return {
        "sentence_length_avg": avg_length,
        "sentence_length_std": std_length,
        "short_sentence_ratio": round(short_count / total, 2),
        "medium_sentence_ratio": round(medium_count / total, 2),
        "long_sentence_ratio": round(long_count / total, 2),
        "question_ratio": round(question_count / total, 2),
        "exclamation_ratio": round(exclamation_count / total, 2),
        "total_sentences": total,
    }


def analyze_paragraph(text: str) -> Dict[str, Any]:
    """
    段落结构分析
    
    Args:
        text: 输入文本
        
    Returns:
        段落特征字典
    """
    paragraphs = _split_paragraphs(text)
    if not paragraphs:
        return {
            "avg_length": 0,
            "length_distribution": {"short": 0, "medium": 0, "long": 0},
        }
    
    lengths = [len(p) for p in paragraphs]
    avg_length = round(statistics.mean(lengths), 1)
    
    short_count = sum(1 for l in lengths if l < 80)
    medium_count = sum(1 for l in lengths if 80 <= l <= 160)
    long_count = sum(1 for l in lengths if l > 160)
    total = len(lengths)
    
    return {
        "avg_length": avg_length,
        "length_distribution": {
            "short": round(short_count / total, 2),
            "medium": round(medium_count / total, 2),
            "long": round(long_count / total, 2),
        },
        "total_paragraphs": total,
    }


def detect_ai_patterns(text: str) -> Dict[str, Any]:
    """
    AI 套话检测
    
    Args:
        text: 输入文本
        
    Returns:
        AI 套话特征字典
    """
    found_patterns = []
    for pattern in AI_PATTERNS:
        matches = re.findall(pattern, text)
        found_patterns.extend(matches)
    
    unique_found = list(set(found_patterns))
    total_patterns = len(AI_PATTERNS)
    detected = len(unique_found)
    detection_rate = round(detected / total_patterns, 2) if total_patterns > 0 else 0
    
    # 生成反 AI 替换规则
    replacements = {
        "首先": ["第一", "先说", "开头来说"],
        "其次": ["第二", "再说"],
        "最后": ["最后说", "末了"],
        "综上所述": ["所以", "总的来看", "说白了"],
        "总而言之": ["所以", "总之"],
        "值得注意的是": ["注意", "有意思的是"],
        "需要指出的是": ["要注意的是", "别忘了"],
        "不可否认": ["确实", "没错"],
        "赋能": ["帮助", "支持"],
        "闭环": ["完整流程", "循环"],
        "抓手": ["方法", "途径"],
        "颗粒度": ["细节程度", "精细度"],
        "对齐": ["统一", "协调"],
        "拉齐": ["统一", "对齐"],
    }
    
    # 只保留文本中实际出现的替换规则
    active_replacements = {}
    for pattern, alternatives in replacements.items():
        if pattern in unique_found or re.search(pattern, text):
            active_replacements[pattern] = alternatives
    
    return {
        "detected_patterns": unique_found,
        "detection_rate": detection_rate,
        "total_patterns_checked": total_patterns,
        "ai_score": round(len(unique_found) / max(len(set(re.findall(r'[\u4e00-\u9fff]{2,}', text))) or 1, 1) * 100, 2),
        "replacements": active_replacements,
    }


def analyze_transitions(text: str) -> Dict[str, Any]:
    """
    过渡词偏好分析
    
    Args:
        text: 输入文本
        
    Returns:
        过渡词特征字典
    """
    personal_transitions = [
        "然后", "不过", "话说回来", "其实", "说白了", "挺", "蛮",
        "反正", "毕竟", "话说", "话说回来", "再说",
    ]
    
    avoided_transitions = [
        "首先", "其次", "最后", "综上所述", "总而言之",
        "一方面", "另一方面", "不仅如此",
    ]
    
    found_personal = [w for w in personal_transitions if w in text]
    found_avoided = [w for w in avoided_transitions if w in text]
    
    return {
        "personal": found_personal,
        "avoided": found_avoided,
    }


def analyze_style(text: str, stop_words: Optional[set] = None) -> Dict[str, Any]:
    """
    综合风格分析（入口函数）
    
    Args:
        text: 输入文本
        stop_words: 自定义停用词
        
    Returns:
        完整风格 Profile 字典
    """
    logger.info(f"开始风格分析，文本长度: {len(text)} 字符")
    
    vocab = analyze_vocabulary(text, stop_words)
    rhythm = analyze_rhythm(text)
    paragraph = analyze_paragraph(text)
    ai_patterns = detect_ai_patterns(text)
    transitions = analyze_transitions(text)
    
    profile = {
        "vocabulary": vocab,
        "rhythm": rhythm,
        "paragraph": paragraph,
        "transitions": transitions,
        "anti_ai": ai_patterns,
    }
    
    logger.info(
        f"风格分析完成: {len(vocab['favorite_words'])} 个标志性词汇, "
        f"平均句长 {rhythm['sentence_length_avg']} 字, "
        f"AI 套话检测率 {ai_patterns['detection_rate']:.0%}"
    )
    
    return profile
